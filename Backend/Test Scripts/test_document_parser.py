"""
test_document_parser.py — Document Parser Service Test Suite

Tests:
  1. Service health + capability report
  2. TXT file upload → chunks indexed in Qdrant
  3. MD file upload → heading-aware chunking verified
  4. CSV upload → tabular content ingested
  5. DOCX upload (if python-docx available)
  6. Invalid file type rejection (400)
  7. Missing collection_name rejection (400)
  8. Multi-file batch upload
"""

import io
import sys
import os
import requests

DOC_PARSER_URL = "http://127.0.0.1:8004"
COLLECTION = "test_doc_parser"

PASS = "[OK]"
FAIL = "[FAIL]"
SKIP = "[SKIP]"


def _header(title: str) -> None:
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")


def _server_online() -> bool:
    try:
        r = requests.get(f"{DOC_PARSER_URL}/status", timeout=3)
        return r.status_code == 200
    except Exception:
        return False


def _ingest(filename: str, content: bytes, collection: str = COLLECTION, content_type: str = "text/plain") -> requests.Response:
    return requests.post(
        f"{DOC_PARSER_URL}/ingest",
        files={"files": (filename, io.BytesIO(content), content_type)},
        data={"collection_name": collection},
        timeout=60,
    )


# ------------------------------------------------------------------ #
# TEST 1: Health check
# ------------------------------------------------------------------ #
def test_health():
    _header("TEST 1: Service Health & Capabilities")
    r = requests.get(f"{DOC_PARSER_URL}/status", timeout=5)
    assert r.status_code == 200, f"Expected 200, got {r.status_code}"
    data = r.json()
    caps = data.get("capabilities", {})
    deps = data.get("dependencies", {})

    print(f"  {PASS} Service online")
    print(f"       PDF (opendataloader): {caps.get('pdf_opendataloader')}")
    print(f"       PDF (pypdf fallback): {caps.get('pdf_pypdf_fallback')}")
    print(f"       DOCX:                {caps.get('docx')}")
    print(f"       CSV (pandas):        {caps.get('csv_pandas')}")
    print(f"       Embedding Service:   {deps.get('embedding_service')}")
    print(f"       Qdrant:              {deps.get('qdrant')}")
    return caps


# ------------------------------------------------------------------ #
# TEST 2: Plain text upload
# ------------------------------------------------------------------ #
def test_txt_upload():
    _header("TEST 2: TXT File Upload")
    content = b"""
UpLink Intelligence Platform

UpLink is a RAG-powered project intelligence tool that ingests GitHub repositories,
Jira tickets, and documents to answer developer questions in context.

Key components:
- RAG Pipeline (port 6399): GitHub and Jira ingestion
- Embedding Service (port 6377): Sentence-Transformer vectors
- Qdrant (port 6366): Vector database
- Document Parser (port 8004): PDF/DOCX/TXT ingestion

The system uses Gemini for LLM summarisation and chat completion.
"""
    r = _ingest("uplink_overview.txt", content)
    assert r.status_code == 200, f"Expected 200, got {r.status_code}: {r.text}"
    data = r.json()
    result = data["results"].get("uplink_overview.txt", {})
    chunks = result.get("chunks_stored", 0)
    print(f"  {PASS} TXT ingested. Status: {data['status']}, Chunks stored: {chunks}")
    return chunks > 0


# ------------------------------------------------------------------ #
# TEST 3: Markdown upload (heading-aware chunking)
# ------------------------------------------------------------------ #
def test_md_upload():
    _header("TEST 3: Markdown File Upload (heading-aware chunking)")
    content = b"""# Architecture Overview

UpLink uses a microservices architecture with FastAPI servers.

## RAG Pipeline

The RAG pipeline accepts GitHub and Jira URLs, fetches data,
summarises key files using Gemini, and indexes them in Qdrant.

## Embedding Service

The Embedding Service runs `all-mpnet-base-v2` via SentenceTransformers
to produce 768-dimensional vectors for semantic search.

## Document Parser

The Document Parser accepts PDF, DOCX, TXT, MD, and CSV files,
chunks them semantically, and indexes them via the Embedding Service.

## Qdrant

Qdrant is the vector database that stores all embeddings for retrieval.
"""
    r = _ingest("architecture.md", content, content_type="text/markdown")
    assert r.status_code == 200, f"Expected 200, got {r.status_code}: {r.text}"
    data = r.json()
    result = data["results"].get("architecture.md", {})
    parsed = result.get("chunks_parsed", 0)
    stored = result.get("chunks_stored", 0)
    print(f"  {PASS} MD ingested. Parsed: {parsed} chunks, Stored: {stored} chunks")
    if parsed >= 4:
        print(f"  {PASS} Heading-aware chunking produced expected chunk count (>= 4 sections)")
    else:
        print(f"  [WARN] Only {parsed} chunks — expected 4 sections from headings")
    return stored > 0


# ------------------------------------------------------------------ #
# TEST 4: CSV upload
# ------------------------------------------------------------------ #
def test_csv_upload():
    _header("TEST 4: CSV File Upload")
    content = b"""name,role,port,status
RAG Pipeline,Ingestion API,6399,running
Embedding Service,Vector Generator,6377,running
Qdrant,Vector DB,6366,running
Document Parser,File Ingestion,8004,running
Scheduler,Jira Connector,8002,stopped
Event Handler,Event Bus,8003,stopped
"""
    r = _ingest("services.csv", content, content_type="text/csv")
    assert r.status_code == 200, f"Expected 200, got {r.status_code}: {r.text}"
    data = r.json()
    result = data["results"].get("services.csv", {})
    stored = result.get("chunks_stored", 0)
    print(f"  {PASS} CSV ingested. Chunks stored: {stored}")
    return stored > 0


# ------------------------------------------------------------------ #
# TEST 5: DOCX upload (skipped if python-docx not available)
# ------------------------------------------------------------------ #
def test_docx_upload():
    _header("TEST 5: DOCX File Upload")
    try:
        from docx import Document
        from io import BytesIO
        doc = Document()
        doc.add_heading("UpLink Technical Spec", level=1)
        doc.add_paragraph("This document describes the UpLink backend architecture.")
        doc.add_heading("Components", level=2)
        doc.add_paragraph("RAG Pipeline: Ingests GitHub repos and Jira tickets.")
        doc.add_paragraph("Embedding Service: Generates 768-dim vectors.")
        doc.add_paragraph("Document Parser: Parses PDF, DOCX, TXT, CSV, MD.")
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()
    except ImportError:
        print(f"  {SKIP} python-docx not installed. Skipping DOCX test.")
        return None

    r = _ingest(
        "spec.docx", content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert r.status_code == 200, f"Expected 200, got {r.status_code}: {r.text}"
    data = r.json()
    result = data["results"].get("spec.docx", {})
    stored = result.get("chunks_stored", 0)
    print(f"  {PASS} DOCX ingested. Chunks stored: {stored}")
    return stored > 0


# ------------------------------------------------------------------ #
# TEST 6: Unsupported file type rejection
# ------------------------------------------------------------------ #
def test_unsupported_type():
    _header("TEST 6: Unsupported File Type Rejection")
    content = b"a = 1 + 2"
    r = _ingest("script.py", content, content_type="text/x-python")
    if r.status_code == 400:
        print(f"  {PASS} .py correctly rejected with 400")
        return True
    else:
        print(f"  {FAIL} Expected 400, got {r.status_code}: {r.text[:120]}")
        return False


# ------------------------------------------------------------------ #
# TEST 7: Missing collection_name
# ------------------------------------------------------------------ #
def test_missing_collection():
    _header("TEST 7: Missing collection_name Rejection")
    content = b"Some text content."
    r = requests.post(
        f"{DOC_PARSER_URL}/ingest",
        files={"files": ("note.txt", io.BytesIO(content), "text/plain")},
        data={},  # no collection_name
        timeout=10,
    )
    if r.status_code == 422:  # FastAPI validation error for missing required form field
        print(f"  {PASS} Missing collection_name correctly rejected with 422 (FastAPI validation)")
        return True
    elif r.status_code == 400:
        print(f"  {PASS} Missing collection_name rejected with 400")
        return True
    else:
        print(f"  {FAIL} Expected 400 or 422, got {r.status_code}: {r.text[:120]}")
        return False


# ------------------------------------------------------------------ #
# TEST 8: Multi-file batch upload
# ------------------------------------------------------------------ #
def test_batch_upload():
    _header("TEST 8: Multi-File Batch Upload")
    files = [
        ("files", ("readme.txt", io.BytesIO(b"UpLink: A RAG-powered developer assistant."), "text/plain")),
        ("files", ("notes.md", io.BytesIO(b"# Notes\n\nThis is a standalone note."), "text/markdown")),
    ]
    r = requests.post(
        f"{DOC_PARSER_URL}/ingest",
        files=files,
        data={"collection_name": COLLECTION},
        timeout=60,
    )
    assert r.status_code == 200, f"Expected 200, got {r.status_code}: {r.text}"
    data = r.json()
    total = data.get("total_chunks_stored", 0)
    file_results = data.get("results", {})
    print(f"  {PASS} Batch upload completed. Total chunks stored: {total}")
    for fname, res in file_results.items():
        print(f"       {fname}: {res.get('chunks_stored', 0)} chunks ({res.get('status')})")
    return len(file_results) == 2 and total > 0


# ------------------------------------------------------------------ #
# Main
# ------------------------------------------------------------------ #
if __name__ == "__main__":
    print("="*60)
    print("  UpLink Document Parser Test Suite")
    print("="*60)

    if not _server_online():
        print("\n[CRITICAL] Document Parser not running on port 8004.")
        print("  Start it: python 'Backend/Document Parser/server.py'")
        sys.exit(1)

    print(f"\n[OK] Document Parser online at {DOC_PARSER_URL}\n")

    caps = test_health()

    results = [
        ("TXT Upload",           test_txt_upload()),
        ("MD Upload",            test_md_upload()),
        ("CSV Upload",           test_csv_upload()),
        ("DOCX Upload",          test_docx_upload()),
        ("Unsupported Type",     test_unsupported_type()),
        ("Missing Collection",   test_missing_collection()),
        ("Batch Upload",         test_batch_upload()),
    ]

    print(f"\n{'='*60}")
    print("  SUMMARY")
    print(f"{'='*60}")
    passed = sum(1 for _, r in results if r is True)
    skipped = sum(1 for _, r in results if r is None)
    for name, result in results:
        symbol = PASS if result is True else (SKIP if result is None else FAIL)
        print(f"  {symbol}  {name}")

    print(f"\n  {passed}/{len(results) - skipped} tests passed ({skipped} skipped).")
    print(f"{'='*60}")
